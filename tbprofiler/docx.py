import sys
from docxtpl import DocxTemplate
from collections import defaultdict
from .models import ProfileResult, VcfQC
from docx import Document
from typing import List
from copy import deepcopy
import logging
from abc import abstractmethod
from .utils import rv2genes, get_gene2drugs
import os



class DocxResultTemplate:
    @abstractmethod
    def write_output(self, result: ProfileResult, conf: dict, outfile: str):
        """Write the output to the Word document"""
        pass

def sanitize(d):
    d = d.replace("-","_")
    return d

def cache_cells(tab) -> List[List[str]]:
    _cells = tab._cells
    numcols = len(tab.columns)
    cells = []
    for row in tab.rows:
        rowcells = []
        for i in range(numcols):
            rowcells.append(_cells.pop(0))
        cells.append(rowcells)
    return cells

def merge_cells(filename: str) -> None:
    doc = Document(filename)
    
    def _merge_cells(tab, rows: List[int], column: int):
        
        if column >= len(tab.columns) - 1:
            return
        if len(rows)==1:
            return 
        c1 = tab.cell(rows[0], column)
        c2 = tab.cell(rows[-1], column)
        c1_font_size = c1.paragraphs[0].runs[0].font.size

        for r in rows[1:]:
            tab.cell(r, column).text = ''
     

        cm = c1.merge(c2)


            
        values_in_next_column = set([tab.rows[r].cells[column+1].text for r in rows])
        for val in values_in_next_column:
            rows_with_val = [r for r in rows if tab.c[r][column+1].text == val]
            _merge_cells(tab, rows_with_val, column+1)
    
    for tab in doc.tables:
        tab.c = cache_cells(tab)
        values_in_next_column = set([tab.rows[r].cells[0].text for r in range(1, len(tab.rows))])
        for val in values_in_next_column:
            rows_with_val = [r for r in range(1, len(tab.rows)) if tab.c[r][0].text == val]

            _merge_cells(tab, rows_with_val, 0)

    doc.save(filename)


class DefaultTemplate(DocxResultTemplate):
    __template_name__ = "default"
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        __dir__ = os.path.dirname(os.path.abspath(__file__))
        # self.template_filename = f"{__dir__}/template.docx"
        self.template_filename = sys.prefix+"/share/tbprofiler/default_template.docx"
    def write_output(
            self, 
            result: ProfileResult, 
            conf: dict,
            output_filename: str):
        
        id2name = rv2genes(conf['bed'])
        db = conf['json_db']
        tier1_genes = set()
        for gene in db:
            for var in db[gene]:
                for ann in db[gene][var]['annotations']:
                    if ann['type']=='drug_resistance':
                        tier1_genes.add(id2name[gene])
        
        tier1_genes.add('mmpL5')

        if isinstance(result.qc, VcfQC):
            qc_check = {d:100 for d in id2name.values()}
            result.notes.append("This report was generated from a VCF where gene coverage is not available. It is assumed all genes are covered well sufficiently.")
        else:
            qc_check = {d.target:d.percent_depth_pass for d in result.qc.target_qc}


        gene2drugs = get_gene2drugs(conf['bed'])
        table = [{'name':d, 'genes':[]} for d in conf['drugs']]
        for gene,drugs in gene2drugs.items():
            if gene not in tier1_genes: continue
            for d in drugs:
                item = [i for i in table if i['name']==d][0]
                item['genes'].append({'name':gene, 'variants':[]})

        for var in result.dr_variants:
            for d in var.drugs:
                item = [i for i in table if i['name']==d['drug']][0]
                gene_item = [i for i in item['genes'] if i['name']==var.gene_name][0]
                gene_item['variants'].append({
                    'change':var.change,
                    'freq':var.freq,
                    'confidence':d['confidence'],
                    'comment':d['comment']
                })

        comments = {}
        rows = []
        for d in table:
            for g in d['genes']:
                if len(g['variants']) == 0:
                    rows.append({
                        'drug':d['name'].replace("_"," ").title(),
                        'gene':g['name'],
                        'qc':qc_check[g['name']],
                        'change':'',
                        'freq':'',
                        'confidence':'',
                        'comment':''
                    })
                else:
                    for v in g['variants']:
                        if v['comment']!="" and v['comment'] not in comments:
                            comments[v['comment']] = len(comments) + 1
                        rows.append({
                            'drug':d['name'].replace("_"," ").title(),
                            'gene':g['name'],
                            'qc':qc_check[g['name']],
                            'change':v['change'],
                            'freq':int(v['freq']*100),
                            'confidence':v['confidence'] if v['confidence']!="" else "Uncertain significance",
                            'comment':comments.get(v['comment'],'')
                        })

        resistant_drugs_tmp = [r['drug'] for r in rows if r['change']!=""]
        resistant_drugs = [d for d in conf['drugs'] if d.title() in resistant_drugs_tmp]
        if len(resistant_drugs)>0:
            result_summary = f"Known resistance variants for {', '.join(resistant_drugs[:-1])} and {resistant_drugs[-1]} detected."
        else:
            result_summary = "No known resistance variants detected."


        poor_coverage_genes = [g for g in qc_check if qc_check[g]<99]
        if len(poor_coverage_genes)>0:
            result.notes.append(f"Insufficient coverage detected in {len(poor_coverage_genes)} genes.")

        fail_variants = []
        fail_comments = {}
        for var in result.qc_fail_variants:
            if not hasattr(var,'drugs'): continue
            for d in var.drugs:

                if d['comment']!="" and d['comment'] not in fail_comments:
                    fail_comments[d['comment']] = len(fail_comments) + 1
                fail_variants.append({
                    'drug':d['drug'],
                    'gene':var.gene_name,
                    'qc':qc_check[var.gene_name],
                    'change':var.change,
                    'confidence':d['confidence'],
                    'freq':int(var.freq*100),
                    'depth':var.depth,
                    'comment':fail_comments.get(d['comment'],'')
                })
            if var.gene_name in ('mmpL5','mmpS5') and var.type=='frameshift_variant':
                for ann in var.annotation:
                    if ann['type']=='who_confidence' and ann['original_mutation']=='LoF':
                        fail_variants.append({
                            'drug':ann['drug'],
                            'gene':var.gene_name,
                            'qc':qc_check[var.gene_name],
                            'change':var.change,
                            'confidence':ann['confidence'],
                            'freq':int(var.freq*100),
                            'depth':var.depth,
                            'comment':fail_comments.get(ann['comment'],'')
                        })
                        
        fail_variants_unique = set([(f['gene'],f['change']) for f in fail_variants])
        if len(fail_variants)>0:
            result.notes.append(f"{len(fail_variants_unique)} resistance variant(s) failed QC checks. These have not been used to generate the mutation report. See QC failed variants table for details.")


        other_variants = []
        other_comments = {}
        for var in result.other_variants:
            drug_annotation = {a['drug']:a for a in var.annotation if a['type']=='who_confidence'}
            for drug_name in var.gene_associated_drugs:
                d = drug_annotation.get(drug_name,{})
                if d.get('comment','')!='' and d['comment'] not in other_comments:
                    other_comments[d['comment']] = len(other_comments) + 1
                other_variants.append({
                    'drug':drug_name,
                    'gene':var.gene_name,
                    'qc':qc_check[var.gene_name],
                    'change':var.change,
                    'confidence':d.get('confidence','Uncertain significance'),
                    'freq':int(var.freq*100),
                    'depth':var.depth,
                    'comment':other_comments.get(d.get('comment',''),'')
                })



        context = {
            'd': result.model_dump(),
            'result_summary': result_summary,
            'rows':rows,
            'comments': comments,
            'qc_check': qc_check,
            'notes': '\n\n'.join(result.notes),
            'fail_variants': fail_variants,
            'fail_comments': fail_comments,
            'other_variants': other_variants,
            'other_comments': other_comments
        }

        tpl = DocxTemplate(self.template_filename)
        tpl.render(context)
        tpl.save(output_filename)

        merge_cells(output_filename)



def write_docx(result: ProfileResult,conf,outfile,template_file = None, plugin = None):
    if template_file is None:
        template_file = sys.prefix+"/share/tbprofiler/default_template.docx"
    
    if plugin:
        plugin_cls = plugin()
        plugin_cls.write_output(result, conf,outfile)

    else:

        output_cls = DefaultTemplate()
        output_cls.write_output(result, conf, outfile)
        # dr_variant_table = []
        # comments = []
        
        # for var in result.dr_variants:
        #     for drug in var.drugs:
        #         mutation = var.change if len(var.change) < 15 else var.change[:11]+"..."+var.change[-3:]
        #         if drug['comment'] not in comments and len(drug['comment']):
        #             comments.append(drug['comment'])
        #         dr_variant_table.append({
        #             'drug': drug['drug'],
        #             'gene': var.gene_name,
        #             'mutation': mutation,
        #             'depth': var.depth,
        #             'frequency': round(var.freq * 100,2),
        #             'confidence': drug['confidence'],
        #             'comment': comments.index(drug['comment'])+1 if len(drug['comment']) else ""
        #         })
        # drugs_found = set([x['drug'] for x in dr_variant_table])
        # for drug in conf['drugs']:
        #     if drug not in drugs_found:
        #         dr_variant_table.append({
        #             'drug': drug,
        #             'gene': "",
        #             'mutation': "",
        #             'depth': "",
        #             'frequency': "",
        #             'confidence': "",
        #             'comment': ""
        #         })

        # gene_qc = []
        # if 'target_qc' in result.qc:
        #     for item in result.qc.target_qc:

        #         gene_qc.append({
        #             'status': 'ok' if item.percent_depth_pass>0.9 else 'fail',
        #             'gene': item.target,
        #             'median_depth': item.median_depth,
        #             'percent_depth_pass': item.percent_depth_pass,
        #         })

        # dr_variant_table = sorted(dr_variant_table,key=lambda x: conf['drugs'].index(x['drug']))

        # other_variants_table = []
        # for var in result.other_variants:
        #     for ann in var.annotation:
        #         other_variants_table.append({
        #             'gene': var.gene_name,
        #             'mutation': var.change,
        #             'depth': var.depth,
        #             'frequency': round(var.freq * 100,2),
        #             'drug': ann['drug'],
        #             'confidence': ann['confidence'],
        #         })
        
        # data = result.model_dump()



        # variables = {
        #     'date':data['timestamp'].strftime("%d %b %Y"),
        #     'sublineage': data['sub_lineage'],
        #     'version': data['pipeline']['software_version'],
        #     'drtype': data['drtype'],
        #     'd': data,
        #     'dr_variants_table': dr_variant_table,
        #     'other_variants_table': other_variants_table,
        #     'comments': comments,
        #     'gene_qc': gene_qc
        # }


        # doc = DocxTemplate(template_file)
        # doc.render(variables)
        # doc.save(outfile)

        # merge_cells(outfile)